import random
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def generate_problems(n=1000):
    problems = []
    seen = set()
    while len(problems) < n:
        op = random.choice(['+', '-'])
        if op == '+':
            s = random.randint(21, 100)
            a = random.randint(1, s - 1)
            b = s - a
        else:
            a = random.randint(21, 100)
            b = random.randint(1, a - 1)
        key = (a, op, b)
        if key in seen:
            continue
        seen.add(key)
        problems.append(f"{a} {op} {b} =")
    return problems

def set_cell_border(cell, top=None, bottom=None, start=None, end=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
        tcPr.append(tcBorders)
    for edge, val in [('top', top), ('bottom', bottom), ('start', start), ('end', end)]:
        if val is not None:
            el = tcBorders.makeelement(qn(f'w:{edge}'), {
                qn('w:val'): val,
                qn('w:sz'): '4',
                qn('w:space'): '0',
                qn('w:color'): '000000',
            })
            existing = tcBorders.find(qn(f'w:{edge}'))
            if existing is not None:
                tcBorders.remove(existing)
            tcBorders.append(el)

def main():
    problems = generate_problems(1000)
    doc = Document()

    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    cols = 5
    rows = (len(problems) + cols - 1) // cols

    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, problem in enumerate(problems):
        r = i // cols
        c = i % cols
        cell = table.cell(r, c)
        p = cell.paragraphs[0]
        run = p.add_run(problem)
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            left_border = 'single' if idx > 0 else 'none'
            right_border = 'single' if idx < cols - 1 else 'none'
            set_cell_border(cell, top='none', bottom='none', start=left_border, end=right_border)

    doc.save('100.docx')
    print("已生成 100.docx")

if __name__ == '__main__':
    main()
